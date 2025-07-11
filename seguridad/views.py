import os
from io import BytesIO
import zipfile
import json
from datetime import datetime
from urllib.parse import quote
from tempfile import NamedTemporaryFile
from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from django.http import FileResponse, JsonResponse
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import padding
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from rest_framework.decorators import api_view
import requests
from docx import Document
from rest_framework.permissions import AllowAny
from rest_framework.decorators import permission_classes

# Ruta base para almacenamiento de claves del sistema
KEY_DIR = "keys"
PRIVATE_KEY_PATH = os.path.join(KEY_DIR, "server_private_key.pem")
PUBLIC_KEY_PATH = os.path.join(KEY_DIR, "server_public_key.pem")

# Inicializar claves si no existen
def generate_key_pair():
    from cryptography.hazmat.primitives.asymmetric import rsa
    private_key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    os.makedirs(KEY_DIR, exist_ok=True)
    with open(PRIVATE_KEY_PATH, "wb") as f:
        f.write(private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption()
        ))
    with open(PUBLIC_KEY_PATH, "wb") as f:
        f.write(private_key.public_key().public_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PublicFormat.SubjectPublicKeyInfo
        ))

def load_public_key():
    with open(PUBLIC_KEY_PATH, "rb") as f:
        return serialization.load_pem_public_key(f.read())

def load_private_key():
    with open(PRIVATE_KEY_PATH, "rb") as f:
        return serialization.load_pem_private_key(f.read(), password=None)

# Firma visible dentro de .docx
def insertar_firma_en_docx(content, metadata):
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(content)
        tmp.flush()
        doc = Document(tmp.name)
        doc.add_paragraph("\n\n---")
        doc.add_paragraph(
            f"📄 Documento firmado digitalmente\n"
            f"👤 Usuario: {metadata['usuario']}\n"
            f"🪪 DNI: {metadata['dni']}\n"
            f"💼 Cargo: {metadata['cargo']}\n"
            f"🕒 Fecha de firma: {metadata['fecha_firma']}"
        )
        tmp_modified = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp_modified.name)
        with open(tmp_modified.name, "rb") as f:
            return f.read()

# Generar claves al iniciar si no existen
generate_key_pair()

@api_view(["GET"])
@permission_classes([AllowAny])
def greynoise_lookup(request, ip):
    try:
        headers = {"key": "YOUR_GREYNOISE_API_KEY"}
        response = requests.get(f"https://api.greynoise.io/v3/community/{ip}", headers=headers)

        if response.status_code == 404:
            return Response({"error": "Not Found"}, status=404)

        if response.status_code == 429:
            return Response({"error": "Too Many Requests"}, status=429)

        response.raise_for_status()
        return Response(response.json())
    except requests.exceptions.RequestException as e:
        return Response({"error": str(e)}, status=500)

class EncryptDocumentView(APIView):
    parser_classes = [MultiPartParser]

    def post(self, request):
        uploaded_file = request.data["file"]
        content = uploaded_file.read()
        original_filename = uploaded_file.name
        base_filename, ext = os.path.splitext(original_filename)

        signer_info = {
            "usuario": "DraftraX",
            "cargo": "Analista de Ciberseguridad",
            "dni": "72809055",
            "fecha_firma": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "documento_original": original_filename
        }

        if ext.lower() == ".docx":
            content = insertar_firma_en_docx(content, signer_info)

        digest = hashes.Hash(hashes.SHA256())
        digest.update(content)
        file_hash = digest.finalize()

        aes_key = os.urandom(32)
        iv = os.urandom(16)
        cipher = Cipher(algorithms.AES(aes_key), modes.CFB(iv))
        encryptor = cipher.encryptor()
        encrypted_content = encryptor.update(content) + encryptor.finalize()

        public_key = load_public_key()
        encrypted_key = public_key.encrypt(
            aes_key,
            padding.OAEP(
                mgf=padding.MGF1(algorithm=hashes.SHA256()),
                algorithm=hashes.SHA256(),
                label=None
            )
        )

        zip_buffer = BytesIO()
        zip_filename = f"{base_filename}_secure.zip"
        quoted_filename = quote(zip_filename)

        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
            zip_file.writestr("encrypted_file.bin", encrypted_content)
            zip_file.writestr("aes_key_encrypted.bin", encrypted_key)
            zip_file.writestr("iv.bin", iv)
            zip_file.writestr(f"original_document{ext}", content)
            zip_file.writestr("signature.sha256", file_hash)
            zip_file.writestr("original_filename.txt", original_filename)
            zip_file.writestr("signature_info.json", json.dumps(signer_info, indent=4))
            zip_file.writestr("public_key.pem", public_key.public_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PublicFormat.SubjectPublicKeyInfo
            ))

            # ⚠️ Solo en entorno de pruebas: incluir clave privada
            with open(PRIVATE_KEY_PATH, "rb") as private_file:
                zip_file.writestr("private_key.pem", private_file.read())

        zip_buffer.seek(0)
        response = FileResponse(zip_buffer, as_attachment=True)
        response["Content-Type"] = "application/zip"
        response["Content-Disposition"] = f"attachment; filename*=UTF-8''{quoted_filename}"
        return response

class VerifySignatureView(APIView):
    parser_classes = [MultiPartParser]

    def post(self, request):
        file = request.data["file"]
        signature = request.data["signature"]

        file_bytes = file.read()
        signature_bytes = signature.read()

        digest = hashes.Hash(hashes.SHA256())
        digest.update(file_bytes)
        file_hash = digest.finalize()

        valid = file_hash == signature_bytes
        return JsonResponse({"valid": valid})

class DecryptDocumentView(APIView):
    parser_classes = [MultiPartParser]

    def post(self, request):
        try:
            uploaded_zip = request.data["zip_file"]
            private_key = load_private_key()

            zip_data = BytesIO(uploaded_zip.read())

            with zipfile.ZipFile(zip_data) as z:
                required_files = ["encrypted_file.bin", "aes_key_encrypted.bin", "iv.bin"]
                for filename in required_files:
                    if filename not in z.namelist():
                        return Response({"error": f"Falta {filename} en el ZIP"}, status=400)

                encrypted_content = z.read("encrypted_file.bin")
                encrypted_key = z.read("aes_key_encrypted.bin")
                iv = z.read("iv.bin")

                original_filename = "decrypted_document"
                if "original_filename.txt" in z.namelist():
                    original_filename = z.read("original_filename.txt").decode()

            aes_key = private_key.decrypt(
                encrypted_key,
                padding.OAEP(
                    mgf=padding.MGF1(algorithm=hashes.SHA256()),
                    algorithm=hashes.SHA256(),
                    label=None
                )
            )

            cipher = Cipher(algorithms.AES(aes_key), modes.CFB(iv))
            decryptor = cipher.decryptor()
            decrypted_content = decryptor.update(encrypted_content) + decryptor.finalize()

            return FileResponse(BytesIO(decrypted_content), as_attachment=True, filename=original_filename)

        except Exception as e:
            return Response({"error": str(e)}, status=500)

@api_view(["GET"])
def get_public_key(request):
    return FileResponse(open("keys/server_public_key.pem", "rb"), as_attachment=True, filename="server_public_key.pem")
